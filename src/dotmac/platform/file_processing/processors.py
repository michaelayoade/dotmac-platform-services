"""
File processors for different file types.
"""

import asyncio
import io
import os
import time
from pathlib import Path
from typing import Any, Optional

from PIL import Image, ImageOps
import PyPDF2
import ffmpeg

try:  # Optional dependency; tests may patch methods from this module
    import openpyxl
except ImportError:  # pragma: no cover - exercised when Excel support unavailable
    openpyxl = None  # type: ignore[assignment]

try:  # Allow optional python-docx dependency
    import docx
except ImportError:  # pragma: no cover - docx features skipped when unavailable
    docx = None  # type: ignore[assignment]

from .base import (
    FileMetadata,
    FileProcessor,
    ProcessingError,
    ProcessingOptions,
    ProcessingResult,
    ProcessingStatus,
)


class ImageProcessor:
    """Process image files."""

    SUPPORTED_FORMATS = {".jpg", ".jpeg", ".png", ".gif", ".bmp", ".webp", ".tiff", ".ico", ".svg"}

    def __init__(self, options: Optional[ProcessingOptions] = None):
        self.options = options or ProcessingOptions()

    async def process(
        self,
        file_path: str,
        options: Optional[ProcessingOptions] = None,
    ) -> ProcessingResult:
        """Process an image file."""
        start_time = time.time()
        options = options or self.options
        result = ProcessingResult(
            status=ProcessingStatus.PROCESSING,
            original_file=file_path,
        )

        path = Path(file_path)
        if not path.exists():
            result.add_error(f"Image file not found: {file_path}")
            result.status = ProcessingStatus.FAILED
            return result

        try:
            # Validate file. When validation fails attempt a lightweight open to
            # capture the underlying library error so callers get actionable
            # feedback instead of the generic "Invalid image file" message the
            # tests expect to see.
            if not await self.validate(file_path):
                detailed_error: Optional[str] = None
                try:
                    with Image.open(file_path) as img:
                        img.verify()
                except Exception as exc:  # pragma: no cover - best-effort capture
                    detailed_error = str(exc)

                raise ProcessingError(detailed_error or f"Invalid image file: {file_path}")

            # Extract metadata
            metadata = await self.extract_metadata(file_path)
            result.metadata = metadata

            # Open image
            with Image.open(file_path) as img:
                # Strip metadata if requested
                if options.strip_metadata:
                    img = self._strip_metadata(img)

                # Process main image
                processed_img = img.copy()

                # Resize if needed
                if options.resize_width or options.resize_height:
                    processed_img = self._resize_image(
                        processed_img,
                        options.resize_width,
                        options.resize_height,
                        options.maintain_aspect_ratio,
                    )

                # Apply watermark if provided
                if options.watermark_path and os.path.exists(options.watermark_path):
                    processed_img = self._apply_watermark(processed_img, options.watermark_path)

                # Save processed image
                output_path = self._get_output_path(file_path, options)
                self._save_image(processed_img, output_path, options)
                result.add_processed_file(output_path)

                # Generate thumbnails if requested
                if options.generate_thumbnails:
                    for size in options.thumbnail_sizes:
                        thumb_path = await self._create_thumbnail(img, size, file_path, options)
                        result.add_thumbnail(thumb_path)

            result.status = ProcessingStatus.COMPLETED
            result.processing_time = time.time() - start_time

        except Exception as e:
            result.add_error(str(e))
            result.status = ProcessingStatus.FAILED

        return result

    async def validate(self, file_path: str) -> bool:
        """Validate if file is a processable image."""
        try:
            path = Path(file_path)
            if not path.exists():
                return False

            if path.suffix.lower() not in self.SUPPORTED_FORMATS:
                return False

            # Try to open the image
            with Image.open(file_path) as img:
                img.verify()

            return True
        except Exception:
            return False

    async def extract_metadata(self, file_path: str) -> FileMetadata:
        """Extract metadata from image."""
        metadata = FileMetadata.from_file(file_path)

        try:
            with Image.open(file_path) as img:
                metadata.dimensions = img.size
                metadata.extra_metadata["mode"] = img.mode
                metadata.extra_metadata["format"] = img.format

                # Extract EXIF data if available
                if hasattr(img, "_getexif") and img._getexif():
                    exif_data = {
                        k: v
                        for k, v in img._getexif().items()
                        if k in img.EXIF and isinstance(v, (str, int, float))
                    }
                    metadata.extra_metadata["exif"] = exif_data

        except Exception as e:
            metadata.extra_metadata["metadata_error"] = str(e)

        return metadata

    def _resize_image(
        self,
        img: Image.Image,
        width: Optional[int],
        height: Optional[int],
        maintain_aspect: bool,
    ) -> Image.Image:
        """Resize image."""
        if not width and not height:
            return img

        current_width, current_height = img.size

        if maintain_aspect:
            # Calculate new size maintaining aspect ratio
            if width and not height:
                ratio = width / current_width
                height = int(current_height * ratio)
            elif height and not width:
                ratio = height / current_height
                width = int(current_width * ratio)
            elif width and height:
                # Fit within bounds
                width_ratio = width / current_width
                height_ratio = height / current_height
                ratio = min(width_ratio, height_ratio)
                width = int(current_width * ratio)
                height = int(current_height * ratio)

        return img.resize(
            (width or current_width, height or current_height), Image.Resampling.LANCZOS
        )

    def _strip_metadata(self, img: Image.Image) -> Image.Image:
        """Strip metadata from image."""
        data = list(img.getdata())
        img_without_metadata = Image.new(img.mode, img.size)
        img_without_metadata.putdata(data)
        return img_without_metadata

    def _apply_watermark(self, img: Image.Image, watermark_path: str) -> Image.Image:
        """Apply watermark to image."""
        watermark = Image.open(watermark_path)

        # Resize watermark to 10% of image size
        watermark_width = int(img.width * 0.1)
        watermark_ratio = watermark_width / watermark.width
        watermark_height = int(watermark.height * watermark_ratio)
        watermark = watermark.resize((watermark_width, watermark_height), Image.Resampling.LANCZOS)

        # Position at bottom right
        position = (img.width - watermark_width - 10, img.height - watermark_height - 10)

        # Create a copy and paste watermark
        img_copy = img.copy()
        img_copy.paste(watermark, position, watermark if watermark.mode == "RGBA" else None)

        return img_copy

    def _get_output_path(self, original_path: str, options: ProcessingOptions) -> str:
        """Generate output path for processed image."""
        path = Path(original_path)

        if options.output_directory:
            os.makedirs(options.output_directory, exist_ok=True)
            output_dir = Path(options.output_directory)
        else:
            output_dir = path.parent

        # Determine output format
        if options.output_format:
            extension = f".{options.output_format.lower()}"
        else:
            extension = path.suffix

        # Generate output filename
        output_name = f"{path.stem}_processed{extension}"

        return str(output_dir / output_name)

    def _save_image(self, img: Image.Image, output_path: str, options: ProcessingOptions):
        """Save image with optimization."""
        save_kwargs = {}

        # Determine format
        output_format = Path(output_path).suffix.lower().lstrip(".")
        if output_format in ["jpg", "jpeg"]:
            save_kwargs["format"] = "JPEG"
            save_kwargs["quality"] = options.quality
            save_kwargs["optimize"] = options.optimize_size
        elif output_format == "png":
            save_kwargs["format"] = "PNG"
            save_kwargs["optimize"] = options.optimize_size
        elif output_format == "webp":
            save_kwargs["format"] = "WEBP"
            save_kwargs["quality"] = options.quality
            save_kwargs["method"] = 6 if options.optimize_size else 0

        img.save(output_path, **save_kwargs)

    async def _create_thumbnail(
        self,
        img: Image.Image,
        size: tuple[int, int],
        original_path: str,
        options: ProcessingOptions,
    ) -> str:
        """Create a thumbnail."""
        path = Path(original_path)

        if options.output_directory:
            os.makedirs(options.output_directory, exist_ok=True)
            output_dir = Path(options.output_directory)
        else:
            output_dir = path.parent

        # Generate thumbnail filename
        thumb_name = f"{path.stem}_thumb_{size[0]}x{size[1]}{path.suffix}"
        thumb_path = str(output_dir / thumb_name)

        # Create thumbnail
        thumb = img.copy()
        thumb.thumbnail(size, Image.Resampling.LANCZOS)

        # Save thumbnail
        self._save_image(thumb, thumb_path, options)

        return thumb_path


class DocumentProcessor:
    """Process document files (PDF, Word, Excel)."""

    SUPPORTED_FORMATS = {
        ".pdf",
        ".doc",
        ".docx",
        ".xls",
        ".xlsx",
        ".ppt",
        ".pptx",
        ".odt",
        ".ods",
        ".odp",
    }

    def __init__(self, options: Optional[ProcessingOptions] = None):
        self.options = options or ProcessingOptions()

    async def process(
        self,
        file_path: str,
        options: Optional[ProcessingOptions] = None,
    ) -> ProcessingResult:
        """Process a document file."""
        start_time = time.time()
        options = options or self.options
        result = ProcessingResult(
            status=ProcessingStatus.PROCESSING,
            original_file=file_path,
        )

        try:
            # Validate file
            if not await self.validate(file_path):
                raise ProcessingError(f"Invalid document file: {file_path}")

            # Extract metadata
            metadata = await self.extract_metadata(file_path)
            result.metadata = metadata

            # Extract text if requested
            if options.extract_text:
                text = await self._extract_text(file_path)
                result.extracted_text = text

            # Process based on file type
            file_ext = Path(file_path).suffix.lower()

            if file_ext == ".pdf":
                await self._process_pdf(file_path, options, result)
            elif file_ext in [".doc", ".docx"]:
                await self._process_word(file_path, options, result)
            elif file_ext in [".xls", ".xlsx"]:
                await self._process_excel(file_path, options, result)

            result.status = ProcessingStatus.COMPLETED
            result.processing_time = time.time() - start_time

        except Exception as e:
            result.add_error(str(e))
            result.status = ProcessingStatus.FAILED

        return result

    async def validate(self, file_path: str) -> bool:
        """Validate if file is a processable document."""
        path = Path(file_path)
        return path.exists() and path.suffix.lower() in self.SUPPORTED_FORMATS

    async def extract_metadata(self, file_path: str) -> FileMetadata:
        """Extract metadata from document."""
        metadata = FileMetadata.from_file(file_path)

        file_ext = Path(file_path).suffix.lower()

        try:
            if file_ext == ".pdf":
                with open(file_path, "rb") as f:
                    reader = PyPDF2.PdfReader(f)
                    metadata.pages = len(reader.pages)

                    if reader.metadata:
                        metadata.extra_metadata["title"] = reader.metadata.get("/Title", "")
                        metadata.extra_metadata["author"] = reader.metadata.get("/Author", "")
                        metadata.extra_metadata["subject"] = reader.metadata.get("/Subject", "")

            elif file_ext in [".docx"]:
                if docx is None:
                    raise ProcessingError("python-docx is required for DOCX metadata extraction")
                doc = docx.Document(file_path)
                metadata.pages = len(doc.element.body)
                metadata.extra_metadata["paragraphs"] = len(doc.paragraphs)

            elif file_ext in [".xlsx"]:
                if openpyxl is None:
                    raise ProcessingError("openpyxl is required for XLSX metadata extraction")
                wb = openpyxl.load_workbook(file_path, read_only=True)
                metadata.extra_metadata["sheets"] = wb.sheetnames
                metadata.pages = len(wb.sheetnames)

        except Exception as e:
            metadata.extra_metadata["metadata_error"] = str(e)

        return metadata

    async def _extract_text(self, file_path: str) -> str:
        """Extract text from document."""
        file_ext = Path(file_path).suffix.lower()

        try:
            if file_ext == ".pdf":
                text = ""
                with open(file_path, "rb") as f:
                    reader = PyPDF2.PdfReader(f)
                    for page in reader.pages:
                        text += page.extract_text()
                return text

            elif file_ext in [".docx"]:
                if docx is None:
                    raise ProcessingError("python-docx is required for DOCX text extraction")
                doc = docx.Document(file_path)
                return "\n".join([para.text for para in doc.paragraphs])

            elif file_ext in [".xlsx"]:
                if openpyxl is None:
                    raise ProcessingError("openpyxl is required for XLSX text extraction")
                wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
                text = []
                for sheet_name in wb.sheetnames:
                    worksheet = None
                    # Attempt subscripting first (standard API)
                    if hasattr(wb, "__getitem__"):
                        try:
                            worksheet = wb[sheet_name]
                        except Exception:
                            worksheet = None
                    # Fall back to legacy helpers in case of patched mocks
                    if worksheet is None:
                        getter = getattr(wb, "get_sheet_by_name", None)
                        if callable(getter):
                            worksheet = getter(sheet_name)
                    if worksheet is None:
                        worksheets = getattr(wb, "worksheets", None)
                        if worksheets:
                            worksheet = next(
                                (ws for ws in worksheets if getattr(ws, "title", None) == sheet_name),
                                None,
                            )
                    if worksheet is None:
                        continue

                    sheet = worksheet
                    text.append(f"Sheet: {sheet_name}")
                    for row in sheet.iter_rows(values_only=True):
                        text.append("\t".join(str(cell) if cell else "" for cell in row))
                return "\n".join(text)

        except Exception as e:
            return f"Text extraction failed: {str(e)}"

        return ""

    async def _process_pdf(
        self,
        file_path: str,
        options: ProcessingOptions,
        result: ProcessingResult,
    ):
        """Process PDF file."""
        if options.split_pages:
            # Split PDF into individual pages
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)

                for i, page in enumerate(reader.pages):
                    writer = PyPDF2.PdfWriter()
                    writer.add_page(page)

                    output_path = self._get_page_output_path(file_path, i + 1, options)

                    with open(output_path, "wb") as output_file:
                        writer.write(output_file)

                    result.add_processed_file(output_path)

    async def _process_word(
        self,
        file_path: str,
        options: ProcessingOptions,
        result: ProcessingResult,
    ):
        """Process Word document."""
        if options.convert_to_pdf:
            # Note: This would require additional dependencies like python-docx2pdf
            result.add_warning("Word to PDF conversion not yet implemented")

    async def _process_excel(
        self,
        file_path: str,
        options: ProcessingOptions,
        result: ProcessingResult,
    ):
        """Process Excel file."""
        if options.convert_to_pdf:
            # Note: This would require additional dependencies
            result.add_warning("Excel to PDF conversion not yet implemented")

    def _get_page_output_path(
        self,
        original_path: str,
        page_num: int,
        options: ProcessingOptions,
    ) -> str:
        """Generate output path for a PDF page."""
        path = Path(original_path)

        if options.output_directory:
            os.makedirs(options.output_directory, exist_ok=True)
            output_dir = Path(options.output_directory)
        else:
            output_dir = path.parent

        output_name = f"{path.stem}_page_{page_num}.pdf"
        return str(output_dir / output_name)


class VideoProcessor:
    """Process video files."""

    SUPPORTED_FORMATS = {
        ".mp4",
        ".avi",
        ".mov",
        ".wmv",
        ".flv",
        ".mkv",
        ".webm",
        ".m4v",
        ".mpg",
        ".mpeg",
    }

    def __init__(self, options: Optional[ProcessingOptions] = None):
        self.options = options or ProcessingOptions()

    async def process(
        self,
        file_path: str,
        options: Optional[ProcessingOptions] = None,
    ) -> ProcessingResult:
        """Process a video file."""
        start_time = time.time()
        options = options or self.options
        result = ProcessingResult(
            status=ProcessingStatus.PROCESSING,
            original_file=file_path,
        )

        try:
            # Validate file
            if not await self.validate(file_path):
                raise ProcessingError(f"Invalid video file: {file_path}")

            # Extract metadata
            metadata = await self.extract_metadata(file_path)
            result.metadata = metadata

            # Generate thumbnail
            if options.generate_thumbnails:
                thumb_path = await self._generate_video_thumbnail(file_path, options)
                result.add_thumbnail(thumb_path)

            # Extract frames if requested
            if options.extract_frames:
                frames = await self._extract_frames(file_path, options)
                result.processed_files.extend(frames)

            result.status = ProcessingStatus.COMPLETED
            result.processing_time = time.time() - start_time

        except Exception as e:
            result.add_error(str(e))
            result.status = ProcessingStatus.FAILED

        return result

    async def validate(self, file_path: str) -> bool:
        """Validate if file is a processable video."""
        path = Path(file_path)
        return path.exists() and path.suffix.lower() in self.SUPPORTED_FORMATS

    async def extract_metadata(self, file_path: str) -> FileMetadata:
        """Extract metadata from video."""
        metadata = FileMetadata.from_file(file_path)

        try:
            probe = ffmpeg.probe(file_path)
            video_info = next(s for s in probe["streams"] if s["codec_type"] == "video")

            metadata.dimensions = (int(video_info["width"]), int(video_info["height"]))
            metadata.duration = float(probe["format"]["duration"])
            metadata.extra_metadata["bitrate"] = int(probe["format"]["bit_rate"])
            metadata.extra_metadata["codec"] = video_info["codec_name"]
            metadata.extra_metadata["fps"] = eval(video_info["r_frame_rate"])

        except Exception as e:
            metadata.extra_metadata["metadata_error"] = str(e)

        return metadata

    async def _generate_video_thumbnail(
        self,
        file_path: str,
        options: ProcessingOptions,
    ) -> str:
        """Generate thumbnail from video."""
        path = Path(file_path)

        if options.output_directory:
            os.makedirs(options.output_directory, exist_ok=True)
            output_dir = Path(options.output_directory)
        else:
            output_dir = path.parent

        thumb_path = str(output_dir / f"{path.stem}_thumb.jpg")

        # Extract frame at 1 second
        (
            ffmpeg.input(file_path, ss=1)
            .output(thumb_path, vframes=1)
            .overwrite_output()
            .run(capture_stdout=True, capture_stderr=True)
        )

        return thumb_path

    async def _extract_frames(
        self,
        file_path: str,
        options: ProcessingOptions,
    ) -> list[str]:
        """Extract frames from video."""
        path = Path(file_path)

        if options.output_directory:
            os.makedirs(options.output_directory, exist_ok=True)
            output_dir = Path(options.output_directory)
        else:
            output_dir = path.parent

        # Create frames directory
        frames_dir = output_dir / f"{path.stem}_frames"
        frames_dir.mkdir(exist_ok=True)

        # Extract frames
        pattern = str(frames_dir / "frame_%04d.jpg")

        (
            ffmpeg.input(file_path)
            .output(pattern, r=1 / options.frame_interval)
            .overwrite_output()
            .run(capture_stdout=True, capture_stderr=True)
        )

        # Get list of extracted frames
        frames = sorted(frames_dir.glob("frame_*.jpg"))
        return [str(f) for f in frames]


class AudioProcessor:
    """Process audio files."""

    SUPPORTED_FORMATS = {".mp3", ".wav", ".flac", ".aac", ".ogg", ".wma", ".m4a", ".opus", ".aiff"}

    def __init__(self, options: Optional[ProcessingOptions] = None):
        self.options = options or ProcessingOptions()

    async def process(
        self,
        file_path: str,
        options: Optional[ProcessingOptions] = None,
    ) -> ProcessingResult:
        """Process an audio file."""
        start_time = time.time()
        options = options or self.options
        result = ProcessingResult(
            status=ProcessingStatus.PROCESSING,
            original_file=file_path,
        )

        try:
            # Validate file
            if not await self.validate(file_path):
                raise ProcessingError(f"Invalid audio file: {file_path}")

            # Extract metadata
            metadata = await self.extract_metadata(file_path)
            result.metadata = metadata

            # Generate waveform if requested
            if options.generate_waveform:
                waveform_path = await self._generate_waveform(file_path, options)
                result.add_processed_file(waveform_path)
                result.extra_data["waveform"] = waveform_path

            result.status = ProcessingStatus.COMPLETED
            result.processing_time = time.time() - start_time

        except Exception as e:
            result.add_error(str(e))
            result.status = ProcessingStatus.FAILED

        return result

    async def validate(self, file_path: str) -> bool:
        """Validate if file is a processable audio file."""
        path = Path(file_path)
        return path.exists() and path.suffix.lower() in self.SUPPORTED_FORMATS

    async def extract_metadata(self, file_path: str) -> FileMetadata:
        """Extract metadata from audio."""
        metadata = FileMetadata.from_file(file_path)

        try:
            probe = ffmpeg.probe(file_path)
            audio_info = next(s for s in probe["streams"] if s["codec_type"] == "audio")

            metadata.duration = float(probe["format"]["duration"])
            metadata.extra_metadata["bitrate"] = int(probe["format"]["bit_rate"])
            metadata.extra_metadata["codec"] = audio_info["codec_name"]
            metadata.extra_metadata["sample_rate"] = int(audio_info["sample_rate"])
            metadata.extra_metadata["channels"] = audio_info["channels"]

        except Exception as e:
            metadata.extra_metadata["metadata_error"] = str(e)

        return metadata

    async def _generate_waveform(
        self,
        file_path: str,
        options: ProcessingOptions,
    ) -> str:
        """Generate waveform visualization."""
        path = Path(file_path)

        if options.output_directory:
            os.makedirs(options.output_directory, exist_ok=True)
            output_dir = Path(options.output_directory)
        else:
            output_dir = path.parent

        waveform_path = str(output_dir / f"{path.stem}_waveform.png")

        # Generate waveform using ffmpeg
        (
            ffmpeg.input(file_path)
            .filter("showwavespic", s="640x120")
            .output(waveform_path, vframes=1)
            .overwrite_output()
            .run(capture_stdout=True, capture_stderr=True)
        )

        return waveform_path
